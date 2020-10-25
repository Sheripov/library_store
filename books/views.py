from django.contrib.auth.mixins import PermissionRequiredMixin
from django.urls import reverse_lazy
from django.views.generic import ListView, DetailView
from django.views.generic.edit import CreateView, UpdateView, DeleteView, FormView

from .forms import FileFieldForm
from .models import Books
from docx import Document


class BookCreate(PermissionRequiredMixin, CreateView):
    permission_required = 'polls.can_add'
    model = Books
    fields = ["title", "author_name", "description", "title_img"]


class BookUpdate(PermissionRequiredMixin, UpdateView):
    permission_required = 'polls.can_edit'
    model = Books
    fields = ["title", "author_name", "description", "title_img"]


class BookDelete(PermissionRequiredMixin, DeleteView):
    permission_required = 'polls.can_delete'
    model = Books
    success_url = reverse_lazy('book-list')


class BookListView(ListView):
    model = Books
    context_object_name = 'book_list'
    queryset = Books.objects.order_by('-id')


class BookDetailView(DetailView):
    model = Books
    context_object_name = 'book_detail'


class FileFieldView(FormView):
    form_class = FileFieldForm
    template_name = 'books/book_upload.html'
    success_url = reverse_lazy('book-list')

    def post(self, request, *args, **kwargs):
        form_class = self.get_form_class()
        form = self.get_form(form_class)
        files = request.FILES.getlist('file_field')
        if form.is_valid():
            for f in files:
                document = Document(f)
                title = document.paragraphs[0].text
                author_name = document.paragraphs[1].text
                description = document.paragraphs[2].text
                title_img = document.paragraphs[3].text
                Books.objects.create(
                    title=title,
                    author_name=author_name,
                    description=description,
                    title_img=title_img
                )
            return self.form_valid(form)
        else:
            return self.form_invalid(form)
